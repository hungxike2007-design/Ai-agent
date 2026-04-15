from flask import Blueprint, render_template, request, jsonify, send_file
import google.generativeai as genai
import pandas as pd
import io
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

ai_bp = Blueprint('ai', __name__)

# Cấu hình AI
genai.configure(api_key="AIzaSyB2DbDIfHl0ac0-TbXWrotaVSYUhLF5DQs")
model = genai.GenerativeModel('gemini-flash-latest')

# Biến tạm lưu dữ liệu và nhận xét AI để dùng cho việc xuất báo cáo
data_cache = {
    "content": "",
    "last_ai_response": ""
}

@ai_bp.route('/dashboard')
def dashboard():
    return render_template('dashboard.html')

@ai_bp.route('/upload', methods=['POST'])
def upload_file():
    file = request.files.get('file')
    user_question = request.form.get('question') 

    if file:
        try:
            df = pd.read_excel(file)
            df = df.fillna("")
            
            data_cache["content"] = df.to_string()
            
            prompt = f"Dữ liệu: {data_cache['content']}\n\nHỏi: {user_question if user_question else 'Tóm tắt bảng này và đưa ra nhận xét chuyên sâu'}"
            response = model.generate_content(prompt)
            
            # Lưu lại câu trả lời của AI để xuất file sau này
            data_cache["last_ai_response"] = response.text
            
            table_html = df.head(10).to_html(classes='data-table', index=False)
            
            return render_template('dashboard.html', 
                                   table_html=table_html, 
                                   ai_response=response.text,
                                   file_name=file.filename)
        except Exception as e:
            return f"Lỗi xử lý file: {e}"
    return "Không có file!"

@ai_bp.route('/ask', methods=['POST'])
def ask():
    user_input = request.json.get("question")
    context = data_cache["content"]
    
    if not context:
        return jsonify({"answer": "Bạn chưa tải file lên!"})
        
    prompt = f"Dựa trên dữ liệu: {context}\n\nTrả lời câu hỏi: {user_input}"
    response = model.generate_content(prompt)
    
    # Cập nhật nội dung chat mới nhất vào cache để nếu khách muốn xuất báo cáo ngay đoạn này
    data_cache["last_ai_response"] = response.text
    return jsonify({"answer": response.text})

# --- TÍNH NĂNG MỚI: XUẤT BÁO CÁO ---

@ai_bp.route('/export/<format>')
def export_report(format):
    content = data_cache.get("last_ai_response", "Chưa có nội dung phân tích.")
    
    if format == 'word':
        doc = Document()
        doc.add_heading('HÙNG STORE - BÁO CÁO PHÂN TÍCH AI', 0)
        doc.add_paragraph("Nội dung phân tích tự động:")
        doc.add_paragraph(content)
        
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return send_file(file_stream, as_attachment=True, download_name="Bao_cao_HungStore.docx")

    elif format == 'pdf':
        file_stream = io.BytesIO()
        c = canvas.Canvas(file_stream, pagesize=letter)
        
        # Lưu ý: PDF mặc định không hỗ trợ tiếng Việt có dấu. 
        # Để chạy demo nhanh, mình dùng font Helvetica (không dấu). 
        # Nếu muốn có dấu, Hùng cần tải file .ttf về và dùng pdfmetrics.registerFont
        c.setFont("Helvetica-Bold", 14)
        c.drawString(100, 750, "HUNG STORE - AI REPORT")
        
        c.setFont("Helvetica", 10)
        text_object = c.beginText(100, 720)
        # Chia nhỏ nội dung để không bị tràn trang
        lines = content.split('\n')
        for line in lines:
            text_object.textLine(line[:100]) # Giới hạn độ dài dòng
        
        c.drawText(text_object)
        c.showPage()
        c.save()
        file_stream.seek(0)
        return send_file(file_stream, as_attachment=True, download_name="Bao_cao_HungStore.pdf")

    return "Định dạng không hỗ trợ"