from flask import Blueprint, render_template, request, jsonify, session, send_file, make_response
import pandas as pd
import io
import os
import uuid
from docx import Document
from services.data_processor import get_cleaning_suggestions, generate_auto_chart, generate_multi_charts, _analyze_dataframe
from services.report_service import get_report_prompt, get_available_styles
from database import get_connection, get_all_system_configs

ai_bp = Blueprint('ai', __name__)

from database import configure_ai, get_key_rotator

def get_model():
    """
    Trả về model hiện tại của rotator.
    """
    return get_key_rotator().get_model()

report_cache = {"last_response": ""}

# ── HÀM XÓA BIỂU ĐỒ TẬP TRUNG ──────────────────────────────────────────────
CHARTS_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'static', 'charts')

def _delete_chart_files(file_ids: list):
    """
    Xóa các file biểu đồ PNG tương ứng với danh sách FileID.
    Mỗi biểu đồ được lưu theo quy ước: static/charts/chart_<file_id>.png
    Hàm bỏ qua lỗi (file không tồn tại / quyền truy cập) một cách im lặng.
    """
    os.makedirs(CHARTS_DIR, exist_ok=True)  # đảm bảo thư mục tồn tại
    deleted, skipped = 0, 0
    for fid in file_ids:
        if not fid:
            continue
        chart_file = os.path.join(CHARTS_DIR, f"chart_{fid}.png")
        if os.path.exists(chart_file):
            try:
                os.remove(chart_file)
                deleted += 1
            except OSError as e:
                print(f"[WARN] Không xóa được biểu đồ chart_{fid}.png: {e}")
                skipped += 1
    return deleted, skipped


def cleanup_orphan_charts():
    """
    Quét thư mục static/charts và xóa mọi file PNG không còn FileID
    tương ứng trong bảng ExcelFiles (biểu đồ mồ côi / tồn dư).
    Trả về số file đã xóa.
    """
    import re
    os.makedirs(CHARTS_DIR, exist_ok=True)
    try:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT FileID FROM ExcelFiles")
        valid_ids = {str(row[0]) for row in cursor.fetchall()}
        conn.close()
    except Exception as e:
        print(f"[WARN] cleanup_orphan_charts – không lấy được danh sách FileID: {e}")
        return 0

    removed = 0
    for fname in os.listdir(CHARTS_DIR):
        m = re.match(r'^chart_(\d+)\.png$', fname)
        if m and m.group(1) not in valid_ids:
            try:
                os.remove(os.path.join(CHARTS_DIR, fname))
                removed += 1
            except OSError:
                pass
    return removed


# ── HÀM LỌC TỪ NGỮ THỪA CỦA AI ────────────────────────────────────────────
import re

# Các cụm mở đầu không cần thiết mà AI thường thêm vào
_FILLER_PATTERNS = [
    r'^(Tất nhiên|Chắc chắn rồi|Được thôi|Rất vui được giúp bạn)[!,.]?\s*',
    r'^(Certainly|Sure|Of course|Absolutely)[!,.]?\s*',
    r'^(Bạn đã hỏi về|Dựa trên dữ liệu bạn cung cấp,?)\s*',
    r'^(Xin chào[!,]?)\s*',
    r'^(Đây là|Dưới đây là)\s*(báo cáo|phân tích|kết quả)[^\n]*:\s*',
    r'\*{3,}',           # xoá *** thừa
    r'^---+\s*$',        # xoá đường kẻ --- thừa ở đầu
]

def clean_ai_response(text: str) -> str:
    """Loại bỏ lời mở đầu dư thừa và ký tự rác từ phản hồi AI."""
    if not text:
        return text
    for pattern in _FILLER_PATTERNS:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE | re.MULTILINE)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Chuyển bullet point dạng * sang - trước khi xóa
    text = re.sub(r'^\s*\*\s+', '- ', text, flags=re.MULTILINE)
    # Xóa toàn bộ dấu * thừa
    text = text.replace('*', '')
    
    return text.strip()


def build_smart_summary(df: 'pd.DataFrame') -> str:
    """
    Tạo bản tóm tắt thống kê cô đọ từ DataFrame.
    Giảm 70-90% số token so với df.to_string().
    """
    import io as _io
    parts = []

    # 1. Cấu trúc cơ bản
    n_rows, n_cols = df.shape
    parts.append(f"Số dòng: {n_rows} | Số cột: {n_cols}")
    parts.append(f"Tên cột: {', '.join(df.columns.tolist())}")

    # 2. Kiểu dữ liệu từng cột (rất ít token)
    dtype_lines = [f"  - {col}: {str(dt)}" for col, dt in df.dtypes.items()]
    parts.append("Kiểu dữ liệu:\n" + "\n".join(dtype_lines))

    # 3. Thống kê mô tả cho cột số (describe)
    num_cols = df.select_dtypes(include='number').columns.tolist()
    if num_cols:
        desc = df[num_cols].describe().round(2)
        buf = _io.StringIO()
        desc.to_string(buf)
        parts.append(f"Thống kê mô tả (cột số):\n{buf.getvalue()}")

    # 4. Giá trị null theo cột
    null_info = df.isnull().sum()
    null_str = ", ".join(f"{c}={v}" for c, v in null_info.items() if v > 0)
    parts.append(f"Giá trị null: {null_str if null_str else 'Không có'}")

    # 5. Cột dạng text: top giá trị xuất hiện nhiều nhất (không gử toàn bộ)
    cat_cols = df.select_dtypes(include='object').columns.tolist()
    for col in cat_cols[:5]:  # tối đa 5 cột text
        top = df[col].value_counts().head(5)
        top_str = ", ".join(f"{k}({v})" for k, v in top.items())
        nunique = df[col].nunique()
        parts.append(f"'{col}': {nunique} giá trị khạc nhau. Top 5: {top_str}")

    # 6. Dữ liệu chi tiết (Gửi toàn bộ để AI có thể xuất hết)
    sample_buf = _io.StringIO()
    df.head(1000).fillna("").to_csv(sample_buf, index=False)
    parts.append(f"Dữ liệu chi tiết:\n{sample_buf.getvalue()}")

    return "\n\n".join(parts)

@ai_bp.route('/dashboard')
def dashboard():
    return render_template('dashboard.html')

@ai_bp.route('/upload', methods=['POST'])
def upload_file():
    file = request.files.get('file')
    style = request.form.get('style_preference', 'Kỹ thuật')
    user_id = session.get('user_id')

    if not file: return "Chưa chọn file!"
    if not user_id: return "Vui lòng đăng nhập lại!"

    try:
        # 1. Đọc và xử lý Excel
        df = pd.read_excel(file, dtype=str)
        filename = file.filename
        
        # Save the file to disk for later retrieval
        import os
        os.makedirs('uploads', exist_ok=True)
        file.seek(0)
        file.save(os.path.join('uploads', filename))
        
        cleaning_hints = get_cleaning_suggestions(df)
        df_display = df.fillna("")

        # Tạo bản tóm tắt thống kê (tiết kiệm 70-90% token so với df.to_string())
        smart_summary = build_smart_summary(df)
        
        # 2. LƯU FILE VÀO DATABASE (Bảng ExcelFiles)
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO ExcelFiles (UserID, FileName, FilePath, UploadDate, Status) 
            OUTPUT INSERTED.FileID 
            VALUES (?, ?, ?, GETDATE(), 'Success')""", 
            (user_id, filename, f"uploads/{filename}"))
        file_id = cursor.fetchone()[0]
        
        # 3. ĐẶT TÊN PHIÊN TỪ TÊN FILE (không gọi AI riêng — tiết kiệm 1 lần API call)
        short_title = filename.rsplit('.', 1)[0][:40]  # dùng tên file làm tiêu đề tạm

        cursor.execute("""
            INSERT INTO ChatSessions (UserID, FileID, StartTime, SessionTitle) 
            OUTPUT INSERTED.SessionID 
            VALUES (?, ?, GETDATE(), ?)""", 
            (user_id, file_id, short_title))
        session_id = cursor.fetchone()[0]
        session['current_session_id'] = session_id 
        session['current_file_id'] = file_id
        conn.commit()

        # 3.5. Phân tích và tạo biểu đồ thông minh (1 biểu đồ duy nhất)
        chart_info = _analyze_dataframe(df)
        chart_type_chosen = chart_info.get('chart_type', 'none')
        chart_reason = chart_info.get('reason', '')
        print(f"[SMART CHART] Chon: {chart_type_chosen} | {chart_reason}")
        chart_path = generate_auto_chart(df, file_id)

        # 4. GỌI AI TẠO BÁO CÁO — dùng smart_summary (tiết kiệm 70-90% token)
        configs = get_all_system_configs()
        default_prompt = configs.get('DefaultPrompt', '').strip()
        base_prompt = get_report_prompt(smart_summary, style)
        if default_prompt:
            # DefaultPrompt từ admin THAY THẾ phần role description, đặt làm system instruction ưu tiên
            # Tách phần dữ liệu khỏi base_prompt để ghép lại
            data_marker = '\n\nDữ liệu cần phân tích:'
            if data_marker in base_prompt:
                data_part = base_prompt[base_prompt.index(data_marker):]
                prompt = f"{default_prompt}\n\n[Quy tắc định dạng: Dùng Markdown (##, **, bảng |col|, -). Bắt đầu ngay nội dung, không viết lời mở đầu thừa.]{data_part}"
            else:
                prompt = f"{default_prompt}\n\n{base_prompt}"
        else:
            prompt = base_prompt

        generation_config = {
            "temperature": configs.get("Temperature", 0.5),
            "max_output_tokens": min(int(configs.get("MaxTokens", 1024)), 1500)
        }

        try:
            import google.generativeai as genai
            _gen_cfg = genai.types.GenerationConfig(
                temperature=configs.get("Temperature", 0.5),
                max_output_tokens=min(int(configs.get("MaxTokens", 1024)), 1500)
            )
            # Dùng rotator.generate() — tự xoay key khi gặp quota
            response = get_key_rotator().generate(prompt, generation_config=_gen_cfg)
            report_content = clean_ai_response(response.text)
        except Exception as ai_err:
            err_str = str(ai_err).lower()
            # Trường hợp TẤT CẢ key đã hết quota (rotator đã thử hết vòng)
            if 'tất cả' in str(ai_err) and 'api key' in err_str:
                report_content = (
                    "## ⚠️ Tất Cả API Key Đã Hết Quota\n\n"
                    f"{ai_err}\n\n"
                    "**Gợi ý:**\n"
                    "- Thêm key mới từ account Google khác vào `GEMINI_API_KEYS` trong `database.py`\n"
                    "- Hoặc đợi đến 07:00 sáng hôm sau để quota tự reset\n"
                    "- Key miễn phí tại: [aistudio.google.com](https://aistudio.google.com)"
                )
            elif 'quota' in err_str or 'resource exhausted' in err_str or '429' in err_str:
                report_content = (
                    "## ⚠️ Hết Hạn Mức API (Quota Exceeded)\n\n"
                    "Hệ thống đã tự động thử xoay vòng qua tất cả key nhưng đều hết quota.\n\n"
                    "**Giải pháp:**\n"
                    "- Thêm key mới vào danh sách `GEMINI_API_KEYS` trong `database.py`\n"
                    "- Key miễn phí: [aistudio.google.com](https://aistudio.google.com) → Get API Key\n"
                    "- Hoặc đợi reset lúc 07:00 sáng"
                )
            elif any(k in err_str for k in ['api_key', 'invalid', 'api key not valid', '400', 'unauthenticated']):
                report_content = (
                    "## 🔑 API Key Không Hợp Lệ\n\n"
                    "Một hoặc nhiều key trong `GEMINI_API_KEYS` bị sai hoặc chưa kích hoạt.\n\n"
                    "**Cách lấy key đúng:**\n"
                    "1. Vào **aistudio.google.com** (đăng nhập Google)\n"
                    "2. Click **Get API Key** → **Create API key in new project**\n"
                    "3. Copy key → dán vào danh sách `GEMINI_API_KEYS` trong `database.py`\n\n"
                    f"> Lỗi gốc: `{ai_err}`"
                )
            elif 'not found' in err_str or 'model' in err_str:
                report_content = (
                    "## 🚫 Lỗi Tên Model\n\n"
                    "Tên model sai. Đảm bảo `GEMINI_MODEL_NAME` trong `database.py` là:\n"
                    "```\nGEMINI_MODEL_NAME = 'gemini-2.0-flash'\n```\n"
                    f"> Lỗi gốc: `{ai_err}`"
                )
            else:
                report_content = f"## ❌ Lỗi AI\n\n```\n{ai_err}\n```"

        report_cache["last_response"] = report_content

        # 5. LƯU NỘI DUNG BÁO CÁO VÀO BẢNG REPORTS
        try:
            summary_text = report_content[:200] + "..." if len(report_content) > 200 else report_content
            sql_report = """
                INSERT INTO Reports (FileID, [Content], CreatedDate, Summary) 
                VALUES (?, ?, GETDATE(), ?)
            """
            cursor.execute(sql_report, (file_id, report_content, summary_text))
            conn.commit()
        except Exception as db_err:
            print(f"Loi khi luu vao bang Reports: {db_err}")

        conn.close()

        return render_template('dashboard.html',
                               table_html=df_display.to_html(classes='table table-hover', index=False),
                               ai_response=report_content,
                               cleaning_hints=cleaning_hints,
                               selected_style=style,
                               chart_path=chart_path,
                               extra_charts=[],
                               chart_type_chosen=chart_type_chosen,
                               chart_reason=chart_reason)
    except Exception as e:
        return f"Loi he thong: {e}"


@ai_bp.route('/ask', methods=['POST'])
def ask():
    data = request.json
    question = data.get('question')
    session_id = session.get('current_session_id') # Lấy ID phiên từ lúc upload
    file_id = session.get('current_file_id')

    # Nếu có session_id mà chưa có file_id, lấy file_id từ DB
    if session_id and not file_id:
        try:
            conn = get_connection()
            cursor = conn.cursor()
            cursor.execute("SELECT FileID FROM ChatSessions WHERE SessionID = ?", (session_id,))
            s_row = cursor.fetchone()
            conn.close()
            if s_row: file_id = s_row[0]
        except: pass

    excel_data = "Không có dữ liệu bảng."
    if file_id:
        try:
            conn = get_connection()
            cursor = conn.cursor()
            cursor.execute("SELECT FilePath FROM ExcelFiles WHERE FileID = ?", (file_id,))
            row = cursor.fetchone()
            conn.close()
            if row and row[0] and os.path.exists(row[0]):
                df = pd.read_excel(row[0], dtype=str)
                excel_data = build_smart_summary(df)
        except Exception as e:
            print("Lỗi đọc file excel cho ask:", e)

    try:
        user_id = int(session.get('user_id'))
    except:
        return jsonify({"answer": "Vui lòng đăng nhập lại."}), 400

    # Nếu chưa có session_id (người dùng chưa upload file mà đã hỏi), tạo mới
    if not session_id:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("INSERT INTO ChatSessions (UserID, StartTime, SessionTitle) OUTPUT INSERTED.SessionID VALUES (?, GETDATE(), ?)", 
                       (user_id, question[:50]))
        session_id = cursor.fetchone()[0]
        session['current_session_id'] = session_id
        conn.commit()
        conn.close()

    try:
        # 1. Tạo tiêu đề bằng AI
        try:
            title_prompt = f"Tóm tắt ngắn gọn câu hỏi này làm tiêu đề lịch sử (max 5 từ): '{question}'"
            title_response = get_key_rotator().generate(title_prompt)
            title = title_response.text.strip().replace("*", "").replace('"', '')
        except:
            title = question[:50]

        import google.generativeai as genai
        configs = get_all_system_configs()
        default_prompt = configs.get('DefaultPrompt', '').strip()
        data_prompt = f"Dữ liệu: {excel_data}\n\nCâu hỏi: {question}\n\nTrả lời ngắn gọn, chính xác bằng tiếng Việt."
        if default_prompt:
            # DefaultPrompt từ admin làm system instruction ưu tiên
            full_prompt = f"{default_prompt}\n\n[Định dạng: Dùng Markdown nếu phù hợp. Bắt đầu thẳng vào câu trả lời.]\n\n{data_prompt}"
        else:
            full_prompt = data_prompt
        
        _gen_cfg = genai.types.GenerationConfig(
            temperature=configs.get("Temperature", 0.7),
            max_output_tokens=int(configs.get("MaxTokens", 2048))
        )
        # Dùng rotator.generate() — tự xoay key khi gặp quota
        response = get_key_rotator().generate(full_prompt, generation_config=_gen_cfg)
        answer = clean_ai_response(response.text)
        
        # 3. LƯU VÀO DATABASE (Cả SessionTitle và ChatMessages)
        conn = get_connection()
        cursor = conn.cursor()
        
        # Cập nhật lại tiêu đề phiên chat cho hay hơn (nếu cần)
        cursor.execute("UPDATE ChatSessions SET SessionTitle = ? WHERE SessionID = ?", (title, session_id))
        
        # Lưu tin nhắn của Người dùng và AI (Quan trọng để hiện lịch sử)
        # Sử dụng [Content] vì đây là từ khóa trong SQL
        cursor.execute("INSERT INTO ChatMessages (SessionID, Role, [Content], CreatedAt) VALUES (?, 'user', ?, GETDATE())", (session_id, question))
        cursor.execute("INSERT INTO ChatMessages (SessionID, Role, [Content], CreatedAt) VALUES (?, 'assistant', ?, GETDATE())", (session_id, answer))
        
        conn.commit()
        conn.close()
        
        return jsonify({"answer": answer})
    except Exception as e:
        return jsonify({"answer": f"Lỗi: {str(e)}"}), 500

@ai_bp.route('/history')
def history():
    user_id = session.get('user_id')
    # Nếu không có user_id trong session (chưa đăng nhập), trả về danh sách trống
    if not user_id: 
        return jsonify([])
    
    try:
        conn = get_connection()
        cursor = conn.cursor()
        
        # CHỐT chặn quan trọng: Chỉ lấy những phiên chat mà UserID khớp với người đang đăng nhập
        sql = """
            SELECT SessionID, SessionTitle 
            FROM ChatSessions 
            WHERE UserID = ? 
            ORDER BY StartTime DESC
        """
        cursor.execute(sql, (user_id,))
        rows = cursor.fetchall()
        conn.close()
        
        return jsonify([{"id": row[0], "title": row[1]} for row in rows])
    except Exception as e:
        print(f"Lỗi history: {e}")
        return jsonify([])

@ai_bp.route('/export_report') # Sửa tên route để hết lỗi 404
def export_report():
    format_type = request.args.get('format', 'word')
    session_id = request.args.get('session_id')
    if not session_id:
        session_id = session.get('current_session_id')
    
    content = ""
    file_id = None

    # 1. Nếu có session_id (từ lịch sử), lấy FileID và nội dung từ Database
    if session_id:
        try:
            conn = get_connection()
            cursor = conn.cursor()
            # Lấy FileID trước
            cursor.execute("SELECT FileID FROM ChatSessions WHERE SessionID = ?", (session_id,))
            s_row = cursor.fetchone()
            if s_row:
                file_id = s_row[0]
                # Lấy Content
                cursor.execute("SELECT [Content] FROM Reports WHERE FileID = ? ORDER BY CreatedDate DESC", (file_id,))
                r_row = cursor.fetchone()
                if r_row:
                    content = r_row[0]
            conn.close()
        except Exception as e:
            print(f"Lỗi DB khi xuất file: {e}")

    # 1.5. Fallback lấy file_id từ session nếu DB không tìm thấy
    if not file_id:
        file_id = session.get('current_file_id')

    # 2. Nếu không tìm thấy nội dung, dùng cache
    if not content:
        content = report_cache.get("last_response", "")

    # 3. Kiểm tra dữ liệu cuối cùng
    if not content: 
        return "Không có dữ liệu để xuất! Hãy chọn một phiên chat có dữ liệu."

    # Làm sạch content: bỏ ký tự NUL và các ký tự điều khiển không in được
    if content:
        content = content.replace('\x00', '')  # NUL character
        content = ''.join(ch for ch in content if ch == '\n' or ch == '\t' or ord(ch) >= 32)

    # Xử lý đường dẫn biểu đồ - normalize để tránh lỗi not found
    chart_path = None
    if file_id:
        # Thử các tên file có thể có
        for candidate in [
            f"static/charts/chart_{file_id}.png",
            f"static/charts/chart_{file_id}.0.png",
        ]:
            full = os.path.join(os.getcwd(), candidate)
            if os.path.exists(full):
                chart_path = candidate
                break
    print(f"[EXPORT] chart_path={chart_path}, file_id={file_id}")

    # --- TIẾN HÀNH XUẤT FILE ---
    from datetime import datetime
    now = datetime.now().strftime("%d/%m/%Y %H:%M")

    if format_type == 'word':
        doc = _build_word_report(content, now, chart_path)
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        return send_file(stream, as_attachment=True, download_name="Bao_cao_AI.docx",
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    # Xuất HTML (in ấn / PDF)
    html_body = _markdown_to_html(content)
    html = _build_html_report(html_body, now, chart_path)
    res = make_response(html)
    res.headers['Content-Disposition'] = 'attachment; filename=Bao_cao_AI.html'
    res.headers['Content-Type'] = 'text/html; charset=utf-8'
    return res


# ── HÀM HỖ TRỢ XUẤT FILE ────────────────────────────────────────────────────

def _markdown_to_html(text: str) -> str:
    """Chuyển đổi Markdown đơn giản thành HTML."""
    import re, html as html_lib
    lines = text.split('\n')
    out = []
    in_table = False
    in_list = False

    for i, raw in enumerate(lines):
        line = raw.rstrip()

        # Bảng Markdown | col | col |
        if line.startswith('|'):
            cells = [c.strip() for c in line.strip('|').split('|')]
            # Dòng phân cách bảng (---|---)
            if all(re.match(r'^-+$', c.replace(':', '')) for c in cells if c):
                continue
            if not in_table:
                out.append('<table class="report-table"><tbody>')
                in_table = True
                tag = 'th'
            else:
                tag = 'td'
            row = ''.join(f'<{tag}>{html_lib.escape(c)}</{tag}>' for c in cells)
            out.append(f'<tr>{row}</tr>')
            continue
        else:
            if in_table:
                out.append('</tbody></table>')
                in_table = False

        # Headings
        m = re.match(r'^(#{1,3})\s+(.+)', line)
        if m:
            if in_list: out.append('</ul>'); in_list = False
            lvl = len(m.group(1)) + 1  # ## → h3, ### → h4
            out.append(f'<h{lvl}>{html_lib.escape(m.group(2))}</h{lvl}>')
            continue

        # Bullet list
        if re.match(r'^[-*]\s+', line):
            if not in_list: out.append('<ul>'); in_list = True
            item = re.sub(r'^[-*]\s+', '', line)
            item = _inline_md(item)
            out.append(f'<li>{item}</li>')
            continue
        else:
            if in_list: out.append('</ul>'); in_list = False

        # HR
        if re.match(r'^---+$', line):
            out.append('<hr>')
            continue

        # Dòng trống
        if not line.strip():
            continue

        # Đoạn văn thông thường
        out.append(f'<p>{_inline_md(html_lib.escape(line))}</p>')

    if in_table: out.append('</tbody></table>')
    if in_list:  out.append('</ul>')
    return '\n'.join(out)


def _inline_md(text: str) -> str:
    """Xử lý inline Markdown: bold, italic, code."""
    import re
    text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
    text = re.sub(r'\*(.+?)\*',     r'<em>\1</em>', text)
    text = re.sub(r'`(.+?)`',       r'<code>\1</code>', text)
    return text


def _build_html_report(html_body: str, timestamp: str, chart_path: str = None) -> str:
    """Tạo trang HTML hoàn chỉnh với CSS đẹp cho in ấn."""
    chart_html = ""
    if chart_path:
        import base64
        import os
        try:
            full_chart_path = os.path.join(os.getcwd(), chart_path)
            if os.path.exists(full_chart_path):
                with open(full_chart_path, "rb") as image_file:
                    encoded_string = base64.b64encode(image_file.read()).decode("utf-8")
                chart_html = f'<div style="text-align:center; margin: 20px 0;"><img src="data:image/png;base64,{encoded_string}" style="max-width: 100%; border-radius: 8px; box-shadow: 0 4px 12px rgba(0,0,0,0.1);"></div>'
        except Exception as e:
            print(f"Lỗi chèn ảnh vào HTML: {e}")

    return f"""<!DOCTYPE html>
<html lang="vi">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Báo cáo phân tích AI</title>
<style>
  @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700;800&display=swap');
  *, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{ font-family: 'Inter', sans-serif; background: #f8fafc; color: #1e293b; line-height: 1.75; }}
  .wrapper {{ max-width: 860px; margin: 40px auto; background: #fff; border-radius: 16px; box-shadow: 0 4px 30px rgba(0,0,0,0.08); overflow: hidden; }}
  .cover {{ background: linear-gradient(135deg, #10a37f 0%, #0d6b57 100%); padding: 40px 48px; color: white; }}
  .cover h1 {{ font-size: 1.9rem; font-weight: 800; margin-bottom: 8px; }}
  .cover .meta {{ font-size: .85rem; opacity: .8; margin-top: 12px; }}
  .cover .meta span {{ margin-right: 20px; }}
  .body {{ padding: 36px 48px; }}
  h2 {{ font-size: 1.15rem; font-weight: 700; color: #0f172a; margin: 1.6em 0 .5em; padding-bottom: 6px; border-bottom: 2px solid #10a37f; }}
  h3 {{ font-size: 1rem; font-weight: 700; color: #1e293b; margin: 1.2em 0 .4em; }}
  h4 {{ font-size: .92rem; font-weight: 700; color: #334155; margin: 1em 0 .3em; }}
  p {{ margin: .5em 0; color: #334155; }}
  ul {{ padding-left: 1.5em; margin: .5em 0; }}
  li {{ margin: .3em 0; color: #334155; }}
  strong {{ color: #0f172a; font-weight: 700; }}
  em {{ color: #64748b; }}
  code {{ background: #f1f5f9; padding: 2px 6px; border-radius: 4px; font-size: .82rem; font-family: monospace; color: #be185d; }}
  hr {{ border: none; border-top: 1px solid #e2e8f0; margin: 1.4em 0; }}
  .report-table {{ width: 100%; border-collapse: collapse; margin: 1em 0; font-size: .88rem; }}
  .report-table th {{ background: #10a37f; color: white; padding: 10px 14px; text-align: left; font-weight: 600; }}
  .report-table td {{ padding: 9px 14px; border-bottom: 1px solid #e2e8f0; color: #334155; }}
  .report-table tr:nth-child(even) td {{ background: #f8fafc; }}
  .footer {{ background: #f8fafc; padding: 16px 48px; border-top: 1px solid #e2e8f0; font-size: .78rem; color: #94a3b8; display: flex; justify-content: space-between; }}
  @media print {{
    body {{ background: white; }}
    .wrapper {{ box-shadow: none; margin: 0; border-radius: 0; }}
    .cover {{ -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
    .report-table th {{ -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
  }}
</style>
</head>
<body>
<div class="wrapper">
  <div class="cover">
    <h1>📊 Báo cáo Phân tích Dữ liệu</h1>
    <p style="opacity:.85;margin-top:6px;">Được tạo tự động bởi AI Agent Analytics</p>
    <div class="meta">
      <span>🕐 {timestamp}</span>
      <span>🤖 Gemini AI</span>
    </div>
  </div>
  <div class="body">
    {chart_html}
    {html_body}
  </div>
  <div class="footer">
    <span>AI Agent Analytics System</span>
    <span>Xuất lúc {timestamp}</span>
  </div>
</div>
</body>
</html>"""


def _build_word_report(content: str, timestamp: str, chart_path: str = None):
    """Tạo file Word có định dạng từ nội dung Markdown và biểu đồ."""
    import re
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Cài đặt trang
    section = doc.sections[0]
    section.page_width  = Inches(8.27)   # A4
    section.page_height = Inches(11.69)
    section.left_margin = section.right_margin = Cm(2.5)
    section.top_margin  = section.bottom_margin = Cm(2.5)

    # ── TIÊU ĐỀ TRANG ────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('BÁO CÁO PHÂN TÍCH DỮ LIỆU')
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x10, 0xA3, 0x7F)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(f'Được tạo bởi AI Agent Analytics  •  {timestamp}')
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()  # khoảng trắng

    # Chèn biểu đồ nếu có
    if chart_path:
        import os
        full_chart_path = os.path.join(os.getcwd(), chart_path)
        if os.path.exists(full_chart_path):
            try:
                p_chart = doc.add_paragraph()
                p_chart.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_chart = p_chart.add_run()
                r_chart.add_picture(full_chart_path, width=Inches(6.0))
                doc.add_paragraph()
            except Exception as e:
                print(f"Lỗi chèn biểu đồ vào docx: {e}")

    # ── PHÂN TÍCH TỪNG DÒNG ──────────────────────────────────────────────────
    in_table_data = []
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        # Bảng Markdown
        if line.startswith('|'):
            # Thu thập tất cả dòng bảng
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            # Lọc dòng phân cách (---|)
            rows = [r for r in table_lines if not re.match(r'^\|[-| :]+\|$', r.strip())]
            if rows:
                cols = [c.strip() for c in rows[0].strip('|').split('|')]
                tbl = doc.add_table(rows=len(rows), cols=len(cols))
                tbl.style = 'Table Grid'
                for ri, row_line in enumerate(rows):
                    cells = [c.strip() for c in row_line.strip('|').split('|')]
                    for ci, cell_text in enumerate(cells):
                        if ci < len(tbl.rows[ri].cells):
                            cell = tbl.rows[ri].cells[ci]
                            cell.text = cell_text
                            if ri == 0:
                                run = cell.paragraphs[0].runs
                                if run: run[0].bold = True
                                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if run else None
                                # Header row: nền xanh
                                tc = cell._tc
                                tcPr = tc.get_or_add_tcPr()
                                shd = OxmlElement('w:shd')
                                shd.set(qn('w:val'), 'clear')
                                shd.set(qn('w:color'), 'auto')
                                shd.set(qn('w:fill'), '10A37F')
                                tcPr.append(shd)
                doc.add_paragraph()
            continue

        # Heading ##
        m = re.match(r'^(#{1,3})\s+(.+)', line)
        if m:
            lvl = len(m.group(1))
            heading_text = re.sub(r'\*\*(.+?)\*\*', r'\1', m.group(2))
            h = doc.add_heading(heading_text, level=min(lvl, 3))
            h.runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            i += 1
            continue

        # Bullet list
        if re.match(r'^[-*]\s+', line):
            item = re.sub(r'^[-*]\s+', '', line)
            item = re.sub(r'\*\*(.+?)\*\*', r'\1', item)  # strip bold markers
            p = doc.add_paragraph(item, style='List Bullet')
            p.runs[0].font.size = Pt(11)
            i += 1
            continue

        # HR
        if re.match(r'^---+$', line):
            i += 1
            continue

        # Dòng trống
        if not line.strip():
            i += 1
            continue

        # Đoạn văn thông thường (xử lý **bold** inline)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        parts = re.split(r'(\*\*.+?\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                run = p.add_run(part)
            run.font.size = Pt(11)

        i += 1

    return doc


# --- CÁC HÀM CŨ ĐỂ TƯƠNG THÍCH (NẾU CẦN) ---
def save_chat_session(user_id, title, file_id=None):
    # Hàm này giờ đã được tích hợp trực tiếp vào luồng /ask để tối ưu hơn
    pass

def get_user_chat_history(user_id):
    # Hàm này đã được tích hợp vào luồng /history
    pass

@ai_bp.route('/get_session/<int:session_id>')
def get_session(session_id):
    user_id = session.get('user_id')
    if not user_id: 
        return jsonify({"error": "Unauthorized"}), 401
    
    try:
        conn = get_connection()
        cursor = conn.cursor()
        
        # 1. Lấy FilePath từ bảng ExcelFiles thông qua ChatSessions
        sql_file = """
            SELECT f.FilePath, f.FileName, s.FileID 
            FROM ChatSessions s
            JOIN ExcelFiles f ON s.FileID = f.FileID
            WHERE s.SessionID = ? AND s.UserID = ?
        """
        cursor.execute(sql_file, (session_id, user_id))
        file_data = cursor.fetchone()

        # 2. Lấy danh sách tin nhắn
        cursor.execute("SELECT Role, [Content] FROM ChatMessages WHERE SessionID = ? ORDER BY CreatedAt ASC", (session_id,))
        messages = [{"role": row[0], "content": row[1]} for row in cursor.fetchall()]

        # 2.5 Lấy báo cáo phân tích ban đầu (từ bảng Reports) và chèn vào đầu danh sách
        if file_data and file_data[2]:
            file_id = file_data[2]
            cursor.execute("SELECT [Content] FROM Reports WHERE FileID = ?", (file_id,))
            report_row = cursor.fetchone()
            if report_row:
                # Chèn báo cáo phân tích ban đầu vào đầu danh sách tin nhắn
                messages.insert(0, {"role": "assistant", "content": report_row[0]})

        # 3. Đọc dữ liệu Excel để hiển thị lại bảng
        table_html = ""
        chart_path = None
        if file_data and os.path.exists(file_data[0]):
            df = pd.read_excel(file_data[0], dtype=str)  # Đọc ALL cột là string để tránh Mã SV bị 2.36e+09
            df = df.fillna("")
            table_html = df.to_html(classes='table table-hover', index=False)
            
            file_id = file_data[2]
            session['current_file_id'] = file_id
            potential_chart = f"static/charts/chart_{file_id}.png"
            if os.path.exists(os.path.join(os.getcwd(), potential_chart)):
                chart_path = "/" + potential_chart
        
        session['current_session_id'] = session_id
        conn.close()

        return jsonify({
            "messages": messages,
            "table_html": table_html,
            "filename": file_data[1] if file_data else "Unknown",
            "chart_path": chart_path
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    
# --- API ĐỔI TÊN PHIÊN CHAT ---
@ai_bp.route('/rename_session/<int:session_id>', methods=['POST'])
def rename_session(session_id):
    data = request.json
    new_title = data.get('new_title')
    user_id = session.get('user_id')

    if not new_title:
        return jsonify({"error": "Tiêu đề không được để trống"}), 400

    try:
        conn = get_connection()
        cursor = conn.cursor()
        # Chỉ cho phép đổi tên nếu phiên chat đó thuộc về user đang đăng nhập
        cursor.execute("UPDATE ChatSessions SET SessionTitle = ? WHERE SessionID = ? AND UserID = ?", 
                       (new_title, session_id, user_id))
        conn.commit()
        conn.close()
        return jsonify({"success": True, "message": "Đã đổi tên thành công"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# --- API CHIA SẺ PHIÊN CHAT ---
@ai_bp.route('/share_session/<int:session_id>', methods=['POST'])
def share_session(session_id):
    user_id = session.get('user_id')
    share_token = str(uuid.uuid4()) # Tạo chuỗi định danh duy nhất

    try:
        conn = get_connection()
        cursor = conn.cursor()
        
        # 1. Lấy FileID từ ChatSession
        cursor.execute("SELECT FileID FROM ChatSessions WHERE SessionID = ? AND UserID = ?", (session_id, user_id))
        row = cursor.fetchone()
        
        if not row or not row[0]:
            return jsonify({"error": "Không tìm thấy báo cáo để chia sẻ"}), 404
        
        file_id = row[0]

        # 2. Cập nhật ShareToken vào bảng Reports (hoặc tạo mới nếu chưa có)
        # Theo ERD của bạn, ShareToken nằm ở bảng Reports
        cursor.execute("UPDATE Reports SET ShareToken = ?, IsPublic = 1 WHERE FileID = ?", (share_token, file_id))
        
        conn.commit()
        conn.close()
        return jsonify({"share_url": f"/ai/view_shared/{share_token}"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# --- API DỌN DẸP BIỂU ĐỒ MỒ CÔI (Admin) ---
@ai_bp.route('/cleanup_charts', methods=['POST'])
def cleanup_charts():
    """
    Quét toàn bộ static/charts và xóa các file PNG mồ côi
    (không còn FileID tương ứng trong bảng ExcelFiles).
    Chỉ Admin mới được gọi API này.
    """
    if session.get('role') != 'Admin':
        return jsonify({"error": "Chỉ Admin mới có quyền thực hiện thao tác này"}), 403

    try:
        removed = cleanup_orphan_charts()
        return jsonify({
            "success": True,
            "message": f"Đã dọn dẹp xong! Xóa {removed} file biểu đồ mồ côi."
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# --- API XÓA SẠCH SESSION ---
@ai_bp.route('/delete_session/<int:session_id>', methods=['DELETE'])
def delete_session(session_id):
    user_id = session.get('user_id')
    
    try:
        conn = get_connection()
        cursor = conn.cursor()

        # 1. Lấy FilePath và FileID trước khi xóa để xử lý file vật lý
        sql_info = """
            SELECT f.FilePath, f.FileID 
            FROM ChatSessions s
            LEFT JOIN ExcelFiles f ON s.FileID = f.FileID
            WHERE s.SessionID = ? AND s.UserID = ?
        """
        cursor.execute(sql_info, (session_id, user_id))
        info = cursor.fetchone()

        # Cho phép xóa ngay cả khi không có file đính kèm
        if info is not None:
            file_path = info[0] if info[0] else None
            file_id = info[1] if info[1] else None

            # 2. Xóa tin nhắn trong ChatMessages (Ràng buộc khóa ngoại)
            cursor.execute("DELETE FROM ChatMessages WHERE SessionID = ?", (session_id,))
            
            if file_id is not None:
                # 3. Xóa báo cáo trong Reports
                cursor.execute("DELETE FROM Reports WHERE FileID = ?", (file_id,))

            # 4. Xóa phiên chat trong ChatSessions
            cursor.execute("DELETE FROM ChatSessions WHERE SessionID = ?", (session_id,))

            if file_id is not None:
                # 5. Xóa file trong ExcelFiles
                cursor.execute("DELETE FROM ExcelFiles WHERE FileID = ?", (file_id,))

            # 6. Xóa file vật lý trên ổ cứng (nếu tồn tại)
            if file_path and os.path.exists(file_path):
                try: os.remove(file_path)
                except: pass
            
            # 7. Xóa file biểu đồ (nếu có)
            if file_id:
                deleted, _ = _delete_chart_files([file_id])
                if deleted:
                    print(f"[INFO] Đã xóa biểu đồ của FileID={file_id}")

            conn.commit()
            return jsonify({"success": True, "message": "Đã xóa sạch toàn bộ dữ liệu liên quan"})
        else:
            return jsonify({"error": "Không tìm thấy phiên chat hoặc bạn không có quyền xóa"}), 404

    except Exception as e:
        if conn: conn.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        if conn: conn.close()

# --- API XÓA NHIỀU SESSION CÙNG LÚC ---
@ai_bp.route('/bulk_delete_sessions', methods=['POST'])
def bulk_delete_sessions():
    user_id = session.get('user_id')
    if not user_id:
        return jsonify({"error": "Bạn cần đăng nhập để thực hiện thao tác này"}), 401
        
    data = request.json
    session_ids = data.get('session_ids', [])
    
    if not session_ids:
        return jsonify({"error": "Không có phiên chat nào được chọn"}), 400

    try:
        conn = get_connection()
        cursor = conn.cursor()
        
        # Chỉ lấy những phiên thuộc về user hiện tại
        placeholders = ','.join(['?'] * len(session_ids))
        sql_info = f"""
            SELECT f.FilePath, f.FileID, s.SessionID 
            FROM ChatSessions s
            LEFT JOIN ExcelFiles f ON s.FileID = f.FileID
            WHERE s.SessionID IN ({placeholders}) AND s.UserID = ?
        """
        params = session_ids + [user_id]
        
        cursor.execute(sql_info, params)
        rows = cursor.fetchall()
        
        if not rows:
            return jsonify({"error": "Không tìm thấy phiên chat hợp lệ hoặc bạn không có quyền xóa"}), 404
            
        file_paths = [r[0] for r in rows if r[0]]
        file_ids = [r[1] for r in rows if r[1]]
        valid_session_ids = [r[2] for r in rows if r[2]]
        
        if valid_session_ids:
            s_placeholders = ','.join(['?'] * len(valid_session_ids))
            
            # Xóa theo thứ tự để không dính khóa ngoại
            cursor.execute(f"DELETE FROM ChatMessages WHERE SessionID IN ({s_placeholders})", valid_session_ids)
            
            if file_ids:
                f_placeholders = ','.join(['?'] * len(file_ids))
                cursor.execute(f"DELETE FROM Reports WHERE FileID IN ({f_placeholders})", file_ids)
            
            cursor.execute(f"DELETE FROM ChatSessions WHERE SessionID IN ({s_placeholders})", valid_session_ids)
            
            if file_ids:
                f_placeholders = ','.join(['?'] * len(file_ids))
                cursor.execute(f"DELETE FROM ExcelFiles WHERE FileID IN ({f_placeholders})", file_ids)
                
            conn.commit()
            
            # Xóa file vật lý (file upload)
            for path in file_paths:
                if path and os.path.exists(path):
                    try: os.remove(path)
                    except: pass
            
            # Xóa file biểu đồ liên quan
            if file_ids:
                deleted, skipped = _delete_chart_files(file_ids)
                print(f"[INFO] Xóa biểu đồ: {deleted} thành công, {skipped} lỗi.")

        return jsonify({"success": True, "message": f"Đã xóa thành công {len(valid_session_ids)} phiên chat"})
        
    except Exception as e:
        if conn: conn.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        if conn: conn.close()

@ai_bp.route('/view_shared/<token>')
def view_shared(token):
    try:
        conn = get_connection()
        cursor = conn.cursor()
        
        # Truy vấn nội dung báo cáo và thông tin file dựa trên ShareToken
        # IsPublic = 1 để đảm bảo báo cáo này đã được người dùng cho phép chia sẻ
        sql = """
            SELECT r.[Content], f.FileName, r.CreatedDate, f.FilePath, f.FileID
            FROM Reports r
            JOIN ExcelFiles f ON r.FileID = f.FileID
            WHERE r.ShareToken = ? AND r.IsPublic = 1
        """
        cursor.execute(sql, (token,))
        row = cursor.fetchone()
        
        if not row:
            conn.close()
            return "<h3>Link chia sẻ không tồn tại hoặc đã bị gỡ bỏ.</h3>", 404

        report_content = row[0]
        filename = row[1]
        created_date = row[2]
        file_path = row[3]
        file_id = row[4]

        # Xử lý đường dẫn biểu đồ
        chart_path = f"/static/charts/chart_{file_id}.png"
        if not os.path.exists(os.path.join(os.getcwd(), chart_path.lstrip("/"))):
            chart_path = None

        # Đọc dữ liệu Excel để hiển thị bảng xem trước (preview) cho người xem
        table_html = ""
        if os.path.exists(file_path):
            try:
                df = pd.read_excel(file_path, dtype=str)
                table_html = df.fillna("").to_html(classes='table table-bordered table-striped', index=False)
            except Exception as e:
                table_html = f"<p>Không thể hiển thị bản xem trước dữ liệu: {e}</p>"

        conn.close()

        # Render ra một template riêng cho người xem
        return render_template('shared_view.html', 
                               content=report_content, 
                               filename=filename, 
                               date=created_date, 
                               table_preview=table_html,
                               chart_path=chart_path)

    except Exception as e:
        return f"Lỗi hệ thống khi tải báo cáo: {str(e)}", 500

# --- API GỬI PHẢN HỒI / ĐÁNH GIÁ ---
@ai_bp.route('/feedback', methods=['POST'])
def submit_feedback():
    """Người dùng gửi phản hồi trải nghiệm về admin."""
    user_id = session.get('user_id')
    if not user_id:
        return jsonify({"error": "Vui lòng đăng nhập!"}), 401

    data = request.json
    rating = data.get('rating')
    comment = data.get('comment', '').strip()
    category = data.get('category', 'Chung')
    session_id = data.get('session_id') or session.get('current_session_id')

    if not rating or not (1 <= int(rating) <= 5):
        return jsonify({"error": "Điểm đánh giá không hợp lệ (1-5 sao)!"}), 400

    from database import save_feedback
    ok = save_feedback(
        user_id=user_id,
        rating=int(rating),
        comment=comment,
        category=category,
        session_id=session_id or None
    )
    if ok:
        return jsonify({"success": True, "message": "Cảm ơn bạn đã gửi phản hồi! 🎉"})
    return jsonify({"error": "Lưu phản hồi thất bại, vui lòng thử lại!"}), 500